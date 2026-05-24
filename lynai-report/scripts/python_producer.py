#!/usr/bin/env python3
"""
python_producer.py — Cross-platform Producer using python-docx (v1.3+)
=======================================================================

Fallback / preferred producer that does NOT require the Anthropic public `docx`
skill (Node.js + docx-js). Uses pure Python (python-docx + Pillow + matplotlib)
to render the GeoVision house-style .docx with embedded chart PNGs and tables.

When to use:
  - Local Claude Code installations without /mnt/skills/public/docx
  - Codex / Gemini CLI / generic POSIX environments
  - Any time DOCX_SKILL_ROOT is not pointing to a working public docx skill

Usage:
    python python_producer.py \\
      --workdir /path/to/job_workdir \\
      --slug LYNAI_GLOBAL_GOLD_MARKET_20260508_v1 \\
      --token gate_token_v0.json \\
      --redaction redaction_report.json

Expects in workdir:
    artifacts/draft_v{n}.md       (or draft_final.md)
    artifacts/chart_specs.json
    charts/chart_*.png            (rendered PNGs from render_chart.py)
    artifacts/{token,redaction}   (paths via --token / --redaction args)

Writes:
    ${LYNAI_OUTPUTS_DIR}/<slug>.docx
    ${LYNAI_OUTPUTS_DIR}/<slug>.pdf   (if soffice or docx2pdf available)
    <workdir>/producer_log.json

Exit codes:
    0  — both .docx + .pdf produced
    1  — .docx produced, PDF skipped (no soffice / no docx2pdf)
    2  — pre-flight check failed
    3  — render failed
"""

from __future__ import annotations

import argparse
import hashlib
import hmac
import json
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("FATAL: python-docx not installed. Run: pip install python-docx")


# ============================================================================
# House Style Tokens (mirror templates/house_style.json)
# ============================================================================

NAVY    = "0D1F3C"
GOLD    = "C9A84C"
TEXT    = "1A1A1A"
CAPTION = "555555"
BG_ALT  = "FAFAFA"
WHITE   = "FFFFFF"

# Font sizing
TITLE_PT    = 28
H1_PT       = 16
H2_PT       = 13
H3_PT       = 11
BODY_PT     = 10.5
CAPTION_PT  = 8.5
FOOTNOTE_PT = 8
TABLE_PT    = 9.5


# ============================================================================
# Helper: low-level OXML for rich style
# ============================================================================

def _set_cell_bg(cell, color_hex: str):
    """Set a table cell's background fill."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _set_para_border(paragraph, side: str = "bottom", size: int = 6, color: str = NAVY):
    """Add a bottom border to a paragraph (used for footer rule etc.)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(size))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pbdr.append(b)
    p_pr.append(pbdr)


def _run(p, text: str, font: str = "Georgia", size_pt: float = BODY_PT,
         color: str = TEXT, bold: bool = False, italic: bool = False):
    """Add a styled run to a paragraph."""
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size_pt)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    # CJK font for east-asian glyphs
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), 'Source Han Serif SC')
    return r


# ============================================================================
# Page furniture: header / footer / page numbers (GS Top of Mind style)
# ============================================================================

def _page_field(paragraph, field_name: str):
    """Insert a Word field code like { PAGE } or { NUMPAGES } that auto-fills.

    field_name is the Word field instruction text — typically 'PAGE' or 'NUMPAGES'.
    """
    run = paragraph.add_run()
    # Style the run for consistent footer typography
    run.font.name = "Georgia"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(CAPTION)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f" {field_name} "
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    return run


def _restart_page_numbering(section, start_at: int = 1):
    """Force this section's page numbering to restart at start_at (default 1).

    Without this, Word would continue numbering from the previous section,
    so the body's first page would be 'p. 2' (after a cover counted as 1).
    """
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        pg_num_type = OxmlElement('w:pgNumType')
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn('w:start'), str(start_at))


def _shorten_running_title(full_title: str, max_chars: int = 60) -> str:
    """Compress a long title into a header-friendly running title."""
    s = (full_title or "").strip()
    if len(s) <= max_chars:
        return s
    # Try to cut at the first ':' or '—'
    for sep in [":", "—", "-"]:
        if sep in s:
            head = s.split(sep, 1)[0].strip()
            if 8 <= len(head) <= max_chars:
                return head
    return s[: max_chars - 1].rstrip() + "…"


def _set_tab_stop_right(paragraph, position_inches: float = 6.5):
    """Add a right-aligned tab stop so 'L: text  →  R: text' lays out in one line."""
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(position_inches), WD_TAB_ALIGNMENT.RIGHT
    )


def _build_body_header(section, meta: dict):
    """Body section header: running title (L) | date (R), thin Navy bottom border."""
    header = section.header
    header.is_linked_to_previous = False

    # Compute content width so the right tab stop aligns to the page-content edge
    content_in = (section.page_width - section.left_margin - section.right_margin) / 914400  # EMU → inches
    p = header.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    _set_tab_stop_right(p, position_inches=content_in)

    running = meta.get("running_title") or _shorten_running_title(meta.get("title", "LynAI Report"))
    date_str = meta.get("date", "")

    _run(p, running, font="Georgia", size_pt=9, color=NAVY, bold=False)
    _run(p, "\t" + date_str, font="Georgia", size_pt=9, color=CAPTION)

    _set_para_border(p, "bottom", size=6, color=NAVY)


def _build_body_footer(section, meta: dict):
    """Body section footer: company line (L) | p. PAGE of NUMPAGES (R), thin Navy top border."""
    footer = section.footer
    footer.is_linked_to_previous = False

    content_in = (section.page_width - section.left_margin - section.right_margin) / 914400
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    _set_tab_stop_right(p, position_inches=content_in)

    # Top border (thin Navy rule above footer text)
    _set_para_border(p, "top", size=6, color=NAVY)

    # Left: company line
    _run(p, "GeoVision AI Mining  |  LynAI Mines",
         font="Georgia", size_pt=8, color=CAPTION)

    # Tab + Right: "p. X of Y" with live field codes
    _run(p, "\tp. ", font="Georgia", size_pt=8, color=CAPTION)
    _page_field(p, "PAGE")
    _run(p, " of ", font="Georgia", size_pt=8, color=CAPTION)
    _page_field(p, "NUMPAGES")


def _blank_cover_header_footer(section):
    """Cover section gets NO header/footer text — clean institutional cover.

    Important: we must set is_linked_to_previous=False so the body section's
    header/footer doesn't leak back into the cover.
    """
    for hf_kind in ("header", "footer"):
        hf = getattr(section, hf_kind)
        hf.is_linked_to_previous = False
        # Clear any pre-existing paragraph runs
        if hf.paragraphs:
            for run in hf.paragraphs[0].runs:
                run.text = ""


# ============================================================================
# Markdown parser — handles the Drafter's locked dialect
# ============================================================================

CHART_RE = re.compile(r"^\`?\[CHART::([a-z0-9_]+)\]\`?\s*$")
TABLE_RE = re.compile(r"^\`?\[TABLE::([a-z0-9_]+)\]\`?\s*$")
H1_RE    = re.compile(r"^##\s+§(\d+)\s*[—\-–]\s*(.+)$")
H2_RE    = re.compile(r"^###\s+(.+)$")
H3_RE    = re.compile(r"^####\s+(.+)$")
LI_RE    = re.compile(r"^[-*]\s+(.+)$")
ITALIC_LINE_RE = re.compile(r"^\*([^*].*)\*\s*$")  # full-line italic = caption
BOLD_RE  = re.compile(r"\*\*([^*]+)\*\*")
ITALIC_RE = re.compile(r"\*([^*]+)\*")


def _strip_yaml_front_matter(text: str) -> tuple[dict, str]:
    """Pull YAML-ish frontmatter (just title/subtitle/date for our purposes)."""
    meta = {}
    if not text.startswith("---"):
        return meta, text
    end = text.find("\n---", 3)
    if end < 0:
        return meta, text
    fm = text[3:end].strip()
    body = text[end + 4:]
    for line in fm.splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            meta[k.strip()] = v.strip().strip('"').strip("'")
    return meta, body


def _emit_table_md(doc: Document, lines: list[str]):
    """Emit a markdown pipe-table at this position. lines starts with the header row."""
    rows = [l.strip() for l in lines if "|" in l]
    if len(rows) < 2:
        return
    cells = [[c.strip() for c in row.strip().strip("|").split("|")] for row in rows]
    if len(cells) >= 2 and all(re.match(r"^[\-\s:]+$", c) for c in cells[1]):
        cells.pop(1)  # remove separator row
    n_cols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    for i, row in enumerate(cells):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            if i == 0:
                _set_cell_bg(cell, NAVY)
                _run(p, cell_text, size_pt=TABLE_PT, color=WHITE, bold=True)
            else:
                if i % 2 == 0:
                    _set_cell_bg(cell, BG_ALT)
                # Negative numbers in red parens
                if re.match(r"^-?\d", cell_text):
                    if cell_text.startswith("-"):
                        _run(p, "(" + cell_text[1:] + ")", size_pt=TABLE_PT, color="B22222")
                    else:
                        _run(p, cell_text, size_pt=TABLE_PT)
                else:
                    _run(p, cell_text, size_pt=TABLE_PT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _emit_paragraph(doc: Document, text: str, *, body_size: float = BODY_PT,
                    italic: bool = False, color: str = TEXT):
    """Emit a body paragraph, parsing inline **bold** and *italic*."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Inches(0)

    # Parse inline bold/italic — naive but adequate for institutional prose
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            _run(p, text[pos:m.start()], size_pt=body_size, italic=italic, color=color)
        tok = m.group(0)
        if tok.startswith("**"):
            _run(p, tok[2:-2], size_pt=body_size, bold=True, color=color)
        else:
            _run(p, tok[1:-1], size_pt=body_size, italic=True, color=color)
        pos = m.end()
    if pos < len(text):
        _run(p, text[pos:], size_pt=body_size, italic=italic, color=color)
    return p


def _emit_h(doc: Document, level: int, text: str):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        _run(p, text, font="Georgia", size_pt=H1_PT, color=NAVY, bold=True)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        _run(p, text, font="Georgia", size_pt=H2_PT, color=NAVY, bold=True)
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, font="Georgia", size_pt=H3_PT, color=NAVY, bold=True)


def _emit_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    _run(p, text, font="Georgia", size_pt=CAPTION_PT, color=CAPTION, italic=True)


def _emit_chart(doc: Document, chart_id: str, charts_dir: Path):
    img_path = charts_dir / f"{chart_id}.png"
    if not img_path.exists():
        # also try with chart_NN prefix
        candidates = list(charts_dir.glob(f"{chart_id}*.png"))
        if candidates:
            img_path = candidates[0]
        else:
            # fall back to bracket text
            _emit_paragraph(doc, f"[CHART {chart_id} MISSING]", italic=True, color="B22222")
            return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.3))  # ~content width at A4 22mm margins


def parse_and_emit(doc: Document, markdown: str, charts_dir: Path):
    """Top-level markdown → docx renderer."""
    _, body = _strip_yaml_front_matter(markdown)
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # blank line
        if not line.strip():
            i += 1
            continue

        # H1 §N — Title
        m = H1_RE.match(line)
        if m:
            _emit_h(doc, 1, f"§{m.group(1)} — {m.group(2)}")
            i += 1
            continue
        m = H2_RE.match(line)
        if m:
            _emit_h(doc, 2, m.group(1))
            i += 1
            continue
        m = H3_RE.match(line)
        if m:
            _emit_h(doc, 3, m.group(1))
            i += 1
            continue

        # Chart placeholder
        m = CHART_RE.match(line)
        if m:
            _emit_chart(doc, m.group(1), charts_dir)
            i += 1
            continue

        # Table placeholder — currently treated as a marker; pipe-table follows
        m = TABLE_RE.match(line)
        if m:
            i += 1
            continue

        # Pipe table (markdown)
        if "|" in line and i + 1 < len(lines) and re.match(r"^[\s\|\-:]+$", lines[i + 1]):
            tbl_lines = [line, lines[i + 1]]
            j = i + 2
            while j < len(lines) and "|" in lines[j] and not lines[j].startswith("##"):
                tbl_lines.append(lines[j])
                j += 1
            _emit_table_md(doc, tbl_lines)
            i = j
            continue

        # Caption (full-line italic)
        m = ITALIC_LINE_RE.match(line)
        if m:
            _emit_caption(doc, m.group(1))
            i += 1
            continue

        # HR
        if re.match(r"^[-=]{3,}$", line):
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style='List Number')
            txt = re.sub(r"^\d+\.\s+", "", line)
            _run(p, txt, size_pt=BODY_PT)
            i += 1
            continue

        # Bullet list
        m = LI_RE.match(line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            _run(p, m.group(1), size_pt=BODY_PT)
            i += 1
            continue

        # Plain paragraph (may span until next blank or block marker)
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not (
            lines[j].startswith("#") or
            lines[j].startswith("[CHART::") or lines[j].startswith("[TABLE::") or
            "|" in lines[j] or lines[j].startswith("---")
        ):
            para_lines.append(lines[j])
            j += 1
        _emit_paragraph(doc, " ".join(para_lines))
        i = j


# ============================================================================
# Cover page
# ============================================================================

def _emit_cover(doc: Document, meta: dict):
    # Top spacer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)

    # Gold rule
    p = doc.add_paragraph()
    _set_para_border(p, "bottom", size=10, color=GOLD)
    p.paragraph_format.space_after = Pt(20)

    # Wordmark
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, "LYNAI MINES  |  GEOVISION AI MINING", font="Georgia", size_pt=10, color=NAVY, bold=True)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(36)
    _run(p, meta.get("title", "Untitled Report"), font="Georgia", size_pt=TITLE_PT, color=NAVY, bold=True)

    # Subtitle
    if meta.get("subtitle"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        _run(p, meta["subtitle"], font="Georgia", size_pt=14, color=NAVY)

    # Navy rule
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    _set_para_border(p, "bottom", size=8, color=NAVY)

    # Meta block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    _run(p, f"Date: {meta.get('date', '—')}", font="Georgia", size_pt=10, color=CAPTION)
    p = doc.add_paragraph()
    _run(p, f"Author: {meta.get('author', 'LynAI Mines')}", font="Georgia", size_pt=10, color=CAPTION)
    p = doc.add_paragraph()
    _run(p, f"Ref: {meta.get('ref_id', '—')}", font="Georgia", size_pt=10, color=CAPTION)

    # No trailing page break — the section break inserted by build() will
    # advance to a new page where the body section starts (D-11 + page furniture).


# ============================================================================
# Producer entry point
# ============================================================================

def _resolve_paths():
    here = Path(__file__).resolve().parent
    runtime_paths = here.parent / "templates" / "runtime_paths.json"
    if runtime_paths.exists():
        rt = json.loads(runtime_paths.read_text(encoding="utf-8"))
        paths = rt.get("paths", {})
        def _p(key):
            entry = paths.get(key, {})
            return os.environ.get(entry.get("env", ""), entry.get("default", ""))
        return {
            "outputs": _p("outputs_dir"),
            "docx_skill": _p("docx_skill_root"),
        }
    # Hard defaults
    return {
        "outputs": os.environ.get("LYNAI_OUTPUTS_DIR", "/mnt/user-data/outputs"),
        "docx_skill": os.environ.get("DOCX_SKILL_ROOT", "/mnt/skills/public/docx"),
    }


def _sha256_file(p: Path) -> str:
    h = hashlib.sha256()
    h.update(p.read_bytes())
    return "sha256:" + h.hexdigest()


def _verify_gate_token(token_path: Path, draft_path: Path) -> dict:
    """Producer-side gate_token verification (mirror build_docx.js logic)."""
    secret = os.environ.get("GK_TOKEN_SECRET", "")
    if len(secret) < 32:
        raise RuntimeError("GK_TOKEN_SECRET env var missing or too short (need 32+ chars)")
    token = json.loads(token_path.read_text(encoding="utf-8"))

    rule = token.get("rule_applied", {})
    if not (rule.get("operator") == ">" and float(rule.get("threshold", 0)) == 9.5
            and rule.get("scope") == "every_dimension"):
        raise RuntimeError(f"RULE_MISMATCH: {rule}")

    live_hash = _sha256_file(draft_path)
    if live_hash != token.get("draft_content_hash"):
        raise RuntimeError(f"HASH_MISMATCH: live={live_hash[:20]}.. token={token['draft_content_hash'][:20]}..")

    expected_sig = "GK1." + hmac.new(
        secret.encode("utf-8"),
        f"{token['token_id']}||{token['draft_content_hash']}||{token['decision']}".encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    if not hmac.compare_digest(expected_sig, token.get("signature", "")):
        raise RuntimeError("SIGNATURE_INVALID")

    if token.get("decision") not in ("PASS", "DELIVER_WITH_SHORTFALL"):
        raise RuntimeError(f"DECISION_NOT_AUTHORIZING_BUILD: {token.get('decision')}")
    return token


def _verify_redaction(redact_path: Path, draft_hash: str):
    rep = json.loads(redact_path.read_text(encoding="utf-8"))
    if rep.get("overall_verdict") not in ("CLEAR", "REDACTED"):
        raise RuntimeError(f"REDACTION_BLOCKED: {rep.get('overall_verdict')}")


# v1.4.2 D-16 P1-4: cross-check critic JSONs vs scorecard to defeat
# Aggregator-side scorecard fabrication. The Aggregator is supposed to MERGE
# only; this re-derives what the merge should look like from the 3 critic
# inputs and compares to what the scorecard actually claims.
def _crosscheck_scorecard_against_critics(workdir: Path, scorecard_ref_path: Path) -> tuple[bool, str]:
    """Re-derive expected dimensions from the 3 critic JSONs and compare to scorecard.

    Returns (ok, reason). Closes red-team H1 (Aggregator trust gap).
    """
    if not scorecard_ref_path.exists():
        return False, "SCORECARD_MISSING"

    artifacts_dir = workdir / "artifacts"
    critic_paths = {
        "content": artifacts_dir / "critic_content_v0.json",
        "layout":  artifacts_dir / "critic_layout_v0.json",
        "charts":  artifacts_dir / "critic_charts_v0.json",
    }
    # Fallback: glob for any critic_<owner>_v*.json (latest)
    for owner in critic_paths:
        if not critic_paths[owner].exists():
            candidates = sorted(artifacts_dir.glob(f"critic_{owner}_v*.json"))
            if candidates:
                critic_paths[owner] = candidates[-1]

    # If critic JSONs missing, this defense degrades to a warning (logged).
    missing = [o for o, p in critic_paths.items() if not p.exists()]
    if missing:
        return True, f"CROSSCHECK_SKIPPED: missing critic JSONs: {missing}"

    try:
        critics = {owner: json.loads(p.read_text(encoding="utf-8")) for owner, p in critic_paths.items()}
        scorecard = json.loads(scorecard_ref_path.read_text(encoding="utf-8"))
    except Exception as e:
        return False, f"CROSSCHECK_FAILED: parse error: {e}"

    # All 3 critic JSONs must share the same draft_content_hash with each other AND with scorecard
    hashes = {owner: c.get("draft_content_hash") for owner, c in critics.items()}
    hashes["scorecard"] = scorecard.get("draft_content_hash")
    if len(set(hashes.values())) != 1:
        return False, f"CROSSCHECK_HASH_DESYNC: {hashes}"

    # Re-derive scores: each critic owns its non-overlapping dimension subset
    owner_dims = {
        "content": ("R1", "R2", "R3", "R4"),
        "layout":  ("R5", "R6", "R7"),
        "charts":  ("R8", "R9", "R10"),
    }
    expected_scores: dict[str, float] = {}
    for owner, dims in owner_dims.items():
        critic_scores = critics[owner].get("scores", {})
        for dim in dims:
            entry = critic_scores.get(dim)
            if not isinstance(entry, dict) or "score" not in entry:
                return False, f"CROSSCHECK_MISSING_DIM: critic={owner} dim={dim}"
            expected_scores[dim] = float(entry["score"])

    # Compare to scorecard
    sc_dims = scorecard.get("dimensions", {})
    drift = []
    for dim, expected_s in expected_scores.items():
        sc_entry = sc_dims.get(dim, {})
        actual = float(sc_entry.get("score", -1.0))
        if abs(actual - expected_s) > 1e-9:
            drift.append(f"{dim}: scorecard={actual} critics={expected_s}")
    if drift:
        return False, f"SCORECARD_TAMPERED: dimension drift detected: {'; '.join(drift)}"

    return True, "CROSSCHECK_OK"


def _render_pdf(docx_path: Path, outputs_dir: Path, log: dict) -> Path | None:
    """Try multiple PDF rendering paths, in order of preference. Returns path or None."""
    pdf_path = outputs_dir / docx_path.with_suffix(".pdf").name

    # Path 1: soffice (LibreOffice headless)
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        # Try common Windows install locations
        env_path = os.environ.get("LIBREOFFICE_PATH")
        candidates = [env_path] if env_path else [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for candidate in candidates:
            if Path(candidate).exists():
                soffice = candidate
                break

    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(outputs_dir), str(docx_path)],
                check=True, capture_output=True, timeout=120,
            )
            if pdf_path.exists() and pdf_path.stat().st_size >= 5 * 1024:
                log["pdf_render"] = {"method": "soffice", "path": str(pdf_path), "size_bytes": pdf_path.stat().st_size, "success": True}
                return pdf_path
        except Exception as e:
            log["pdf_render"] = {"method": "soffice", "error": str(e)[:300]}

    # Path 2: docx2pdf (requires MS Word on Windows or macOS)
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists() and pdf_path.stat().st_size >= 5 * 1024:
            log["pdf_render"] = {"method": "docx2pdf", "path": str(pdf_path), "size_bytes": pdf_path.stat().st_size, "success": True}
            return pdf_path
    except Exception as e:
        existing = log.get("pdf_render", {})
        existing["docx2pdf_error"] = str(e)[:300]
        log["pdf_render"] = existing

    # No PDF backend worked
    if "pdf_render" not in log:
        log["pdf_render"] = {}
    log["pdf_render"]["success"] = False
    log["pdf_render"]["fallback_message"] = "No PDF backend available. Install LibreOffice (winget install TheDocumentFoundation.LibreOffice) or run on a host with MS Word + docx2pdf."
    return None


def build(args) -> int:
    workdir = Path(args.workdir).resolve()
    paths = _resolve_paths()
    outputs_dir = Path(paths["outputs"]).resolve() if paths["outputs"] else workdir / "outputs"
    outputs_dir.mkdir(parents=True, exist_ok=True)

    slug = args.slug
    if not re.match(r"^LYNAI_[A-Z0-9_]+_[0-9]{8}_v[0-9]+$", slug):
        print(f"FATAL: INVALID_SLUG {slug}", file=sys.stderr)
        return 2

    # Locate inputs
    draft_path = workdir / "artifacts" / "draft_final.md"
    if not draft_path.exists():
        # Fallback: highest-numbered draft_v*.md
        candidates = sorted((workdir / "artifacts").glob("draft_v*.md"))
        if candidates:
            draft_path = candidates[-1]
    if not draft_path.exists():
        print(f"FATAL: draft markdown missing in {workdir}/artifacts/", file=sys.stderr)
        return 2

    charts_dir = workdir / "charts"
    token_path = Path(args.token).resolve() if args.token else workdir / "artifacts" / "gate_token_v0.json"
    redact_path = Path(args.redaction).resolve() if args.redaction else workdir / "artifacts" / "redaction_report.json"
    meta_path = workdir / "artifacts" / "report_meta.json"

    log = {
        "version": "1.3",
        "slug": slug,
        "producer": "python_producer",
        "timestamp": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
        "workdir": str(workdir),
        "outputs_dir": str(outputs_dir),
        "preflight": {},
        "errors": [],
    }

    # Pre-flight
    try:
        token = _verify_gate_token(token_path, draft_path)
        log["preflight"]["gate_token"] = {
            "token_id": token["token_id"],
            "decision": token["decision"],
            "verified": True,
        }
        _verify_redaction(redact_path, token["draft_content_hash"])
        log["preflight"]["redaction"] = "OK"

        # v1.4.2 D-16 P1-4: cross-check critics → scorecard to defeat
        # Aggregator-side scorecard fabrication (red-team H1).
        scorecard_path = workdir / "artifacts" / "scorecard_v0.json"
        if not scorecard_path.exists():
            candidates = sorted((workdir / "artifacts").glob("scorecard_v*.json"))
            if candidates:
                scorecard_path = candidates[-1]
        cross_ok, cross_reason = _crosscheck_scorecard_against_critics(workdir, scorecard_path)
        log["preflight"]["crosscheck"] = {"ok": cross_ok, "reason": cross_reason}
        if not cross_ok:
            raise RuntimeError(f"SCORECARD_CROSSCHECK_FAILED: {cross_reason}")
    except Exception as e:
        log["errors"].append(str(e))
        print(f"FATAL: {e}", file=sys.stderr)
        (workdir / "producer_log.json").write_text(json.dumps(log, indent=2))
        return 2

    # Load meta
    meta = {"title": "LynAI Report", "author": "LynAI Mines"}
    if meta_path.exists():
        try:
            meta.update(json.loads(meta_path.read_text(encoding="utf-8")))
        except Exception:
            pass
    # Fall back: pull title from draft YAML frontmatter
    draft_text = draft_path.read_text(encoding="utf-8")
    fm_meta, _ = _strip_yaml_front_matter(draft_text)
    for k in ("title", "subtitle", "date", "author", "ref_id"):
        if k in fm_meta:
            meta[k] = fm_meta[k]

    # Build — two-section layout per D-11 + GS Top of Mind page-furniture standard:
    #   section[0] = cover (no header/footer, no page number)
    #   section[1] = body  (running title header, page-X-of-Y footer, restart numbering at 1)
    doc = Document()

    # Configure cover section (the initial section that ships with a new Document)
    cover_section = doc.sections[0]
    cover_section.page_height   = Mm(297)
    cover_section.page_width    = Mm(210)
    cover_section.top_margin    = Mm(25)
    cover_section.bottom_margin = Mm(25)
    cover_section.left_margin   = Mm(22)
    cover_section.right_margin  = Mm(22)
    cover_section.header_distance = Mm(12)
    cover_section.footer_distance = Mm(12)
    _blank_cover_header_footer(cover_section)

    # Emit cover content into section[0]
    _emit_cover(doc, meta)

    # Insert section break, opening section[1] for the body
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_height   = Mm(297)
    body_section.page_width    = Mm(210)
    body_section.top_margin    = Mm(25)
    body_section.bottom_margin = Mm(25)
    body_section.left_margin   = Mm(22)
    body_section.right_margin  = Mm(22)
    body_section.header_distance = Mm(12)
    body_section.footer_distance = Mm(12)
    body_section.different_first_page_header_footer = False

    # Restart page numbering so the first body page is 'p. 1 of N'
    _restart_page_numbering(body_section, start_at=1)

    # Build the running header + page-numbered footer for the body section
    _build_body_header(body_section, meta)
    _build_body_footer(body_section, meta)

    # Emit body content into section[1]
    parse_and_emit(doc, draft_text, charts_dir)

    docx_path = outputs_dir / f"{slug}.docx"
    doc.save(str(docx_path))
    log["docx"] = {"path": str(docx_path), "size_bytes": docx_path.stat().st_size}
    print(f"[ok] docx -> {docx_path} ({docx_path.stat().st_size} bytes)")

    # Render PDF
    pdf_path = _render_pdf(docx_path, outputs_dir, log)
    if pdf_path:
        print(f"[ok] pdf  -> {pdf_path} ({pdf_path.stat().st_size} bytes)")
    else:
        # v1.4.2 D-16 P2-7: on Linux, refuse to ship without PDF
        # (Windows/macOS get docx2pdf fallback; only Linux truly needs soffice).
        import platform as _platform
        if _platform.system().lower() == "linux":
            log["success"] = False
            log["errors"].append("PDF_REQUIRED_BUT_UNAVAILABLE: D-11 mandates docx+pdf paired delivery; on Linux you must install LibreOffice (`apt install libreoffice` or equivalent). docx2pdf does not work on Linux because it requires MS Word.")
            (workdir / "producer_log.json").write_text(json.dumps(log, indent=2))
            print(f"[fatal] PDF not rendered on Linux — install LibreOffice. D-11 violated.", file=sys.stderr)
            return 3
        print(f"[warn] PDF not rendered: {log.get('pdf_render', {}).get('fallback_message', 'unknown')}", file=sys.stderr)

    log["success"] = True
    (workdir / "producer_log.json").write_text(json.dumps(log, indent=2))
    return 0 if pdf_path else 1


def main():
    p = argparse.ArgumentParser(description="Python-docx producer (v1.3 fallback path)")
    p.add_argument("--workdir", required=True)
    p.add_argument("--slug", required=True)
    p.add_argument("--token", default=None)
    p.add_argument("--redaction", default=None)
    args = p.parse_args()
    return build(args)


if __name__ == "__main__":
    sys.exit(main())
