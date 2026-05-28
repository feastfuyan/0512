#!/usr/bin/env python3
"""
report_parser.py — 报告文件统一解析入口

职责：读取 PDF/Word/Markdown/HTML 文件，提取文本内容供评分引擎使用。
不依赖外部 LLM 库，纯本地解析。
"""

import os
from pathlib import Path


def parse_report(file_path: str) -> dict:
    """
    统一报告解析入口。

    Args:
        file_path: 报告文件路径（支持 .pdf/.docx/.doc/.md/.html）

    Returns:
        dict: {
            "text": str,           # 提取的纯文本
            "tables": list,        # 表格数据（如有）
            "char_count": int,     # 字符数
            "format": str,         # 文件格式
            "error": str | None    # 错误信息
        }
    """
    ext = Path(file_path).suffix.lower()

    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".doc": _parse_docx,
        ".md": _parse_markdown,
        ".markdown": _parse_markdown,
        ".html": _parse_html,
        ".htm": _parse_html,
    }

    parser = parsers.get(ext)
    if not parser:
        return {
            "text": "",
            "tables": [],
            "char_count": 0,
            "format": ext,
            "error": f"不支持的文件格式: {ext}"
        }

    try:
        return parser(file_path)
    except Exception as e:
        return {
            "text": "",
            "tables": [],
            "char_count": 0,
            "format": ext,
            "error": str(e)
        }


def _parse_pdf(file_path: str) -> dict:
    """解析 PDF 文件"""
    text_parts = []

    # 优先用 PyMuPDF（fitz），回退到 pdfminer
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
    except ImportError:
        try:
            from pdfminer.high_level import extract_text
            text_parts.append(extract_text(file_path))
        except ImportError:
            # 最后回退：用系统命令
            import subprocess
            result = subprocess.run(
                ["pdftotext", "-layout", file_path, "-"],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                text_parts.append(result.stdout)
            else:
                return {
                    "text": "",
                    "tables": [],
                    "char_count": 0,
                    "format": ".pdf",
                    "error": "无法解析PDF：需要安装 PyMuPDF、pdfminer 或 pdftotext"
                }

    text = "\n".join(text_parts)
    return {
        "text": text,
        "tables": [],
        "char_count": len(text),
        "format": ".pdf",
        "error": None
    }


def _parse_docx(file_path: str) -> dict:
    """解析 Word 文件"""
    try:
        from docx import Document
        doc = Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)

        text = "\n".join(paragraphs)
        return {
            "text": text,
            "tables": tables,
            "char_count": len(text),
            "format": ".docx",
            "error": None
        }
    except ImportError:
        return {
            "text": "",
            "tables": [],
            "char_count": 0,
            "format": ".docx",
            "error": "无法解析Word：需要安装 python-docx"
        }


def _parse_markdown(file_path: str) -> dict:
    """解析 Markdown 文件"""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    return {
        "text": text,
        "tables": [],
        "char_count": len(text),
        "format": ".md",
        "error": None
    }


def _parse_html(file_path: str) -> dict:
    """解析 HTML 文件，提取可见文本"""
    with open(file_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    # 简单提取文本（去标签）
    import re
    # 去掉 script/style 标签及内容
    html_clean = re.sub(r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL)
    html_clean = re.sub(r"<style[^>]*>.*?</style>", "", html_clean, flags=re.DOTALL)
    # 去掉所有 HTML 标签
    text = re.sub(r"<[^>]+>", " ", html_clean)
    # 清理多余空白
    text = re.sub(r"\s+", " ", text).strip()
    # 处理 HTML 实体
    import html as html_mod
    text = html_mod.unescape(text)

    return {
        "text": text,
        "tables": [],
        "char_count": len(text),
        "format": ".html",
        "error": None
    }


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python report_parser.py <file_path>")
        sys.exit(1)

    result = parse_report(sys.argv[1])
    if result["error"]:
        print(f"Error: {result['error']}")
    else:
        print(f"Format: {result['format']}")
        print(f"Chars: {result['char_count']}")
        print(f"Tables: {len(result['tables'])}")
        print(f"Preview: {result['text'][:200]}...")
